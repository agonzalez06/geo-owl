#!/usr/bin/env python3
"""
ANC Sheet Generator
Automatically generates daily Admission Number Control sheets by pulling
schedule data from Amion and applying admission order logic.

Configuration is loaded from anc_config.yaml in the same directory.
Edit that file to change admission patterns, add teams, or modify settings.
"""

import urllib.request
import urllib.parse
from datetime import datetime, timedelta
from dataclasses import dataclass, field
from typing import Optional
import re
import os
import subprocess
import tempfile
import logging
import time
import functools
from pathlib import Path

import yaml

# =============================================================================
# LOGGING SETUP
# =============================================================================

def setup_logging() -> logging.Logger:
    """Set up logging to file and console."""
    log_dir = Path(__file__).parent / "logs"
    log_dir.mkdir(exist_ok=True)

    # Log file with date
    log_file = log_dir / f"anc_generator_{datetime.now().strftime('%Y-%m-%d')}.log"

    # Create logger
    logger = logging.getLogger('anc_generator')
    logger.setLevel(logging.DEBUG)

    # File handler - detailed logs
    file_handler = logging.FileHandler(log_file)
    file_handler.setLevel(logging.DEBUG)
    file_formatter = logging.Formatter(
        '%(asctime)s - %(levelname)s - %(funcName)s - %(message)s'
    )
    file_handler.setFormatter(file_formatter)

    # Console handler - less verbose
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_formatter = logging.Formatter('%(levelname)s: %(message)s')
    console_handler.setFormatter(console_formatter)

    # Only add handlers if not already added
    if not logger.handlers:
        logger.addHandler(file_handler)
        logger.addHandler(console_handler)

    return logger

# Initialize logger
logger = setup_logging()

# =============================================================================
# RETRY DECORATOR FOR API CALLS
# =============================================================================

def retry_on_failure(max_retries: int = 3, delay: float = 2.0, backoff: float = 2.0):
    """
    Decorator that retries a function on failure with exponential backoff.

    Args:
        max_retries: Maximum number of retry attempts
        delay: Initial delay between retries (seconds)
        backoff: Multiplier for delay after each retry
    """
    def decorator(func):
        @functools.wraps(func)
        def wrapper(*args, **kwargs):
            current_delay = delay
            last_exception = None

            for attempt in range(max_retries + 1):
                try:
                    return func(*args, **kwargs)
                except Exception as e:
                    last_exception = e
                    if attempt < max_retries:
                        logger.warning(
                            f"{func.__name__} failed (attempt {attempt + 1}/{max_retries + 1}): {e}. "
                            f"Retrying in {current_delay:.1f}s..."
                        )
                        time.sleep(current_delay)
                        current_delay *= backoff
                    else:
                        logger.error(
                            f"{func.__name__} failed after {max_retries + 1} attempts: {e}"
                        )

            raise last_exception
        return wrapper
    return decorator

# =============================================================================
# FAILURE NOTIFICATION
# =============================================================================

def send_failure_notification(date: datetime, error: str, config: dict = None) -> bool:
    """
    Send an email notification when ANC generation fails.

    Args:
        date: The date that failed to generate
        error: Error message describing the failure
        config: Optional config dict (uses ANC_CONFIG if not provided)

    Returns:
        True if notification sent successfully
    """
    config = config or ANC_CONFIG
    notification_config = config.get('notifications', {})

    if not notification_config.get('enabled', True):
        logger.debug("Failure notifications are disabled")
        return False

    recipients = notification_config.get('failure_recipients', [])
    if not recipients:
        # Fall back to email config in email_utils
        try:
            import json
            email_config_path = Path(__file__).parent / "email_utils" / "config.json"
            if email_config_path.exists():
                with open(email_config_path) as f:
                    email_config = json.load(f)
                sender = email_config.get('email', {}).get('sender_email')
                if sender:
                    recipients = [sender]
        except Exception as e:
            logger.warning(f"Could not load email config for failure notification: {e}")

    if not recipients:
        logger.warning("No recipients configured for failure notification")
        return False

    subject = f"ANC Generation Failed - {date.strftime('%A, %B %d, %Y')}"

    # Get the log file path for reference
    log_dir = Path(__file__).parent / "logs"
    log_file = log_dir / f"anc_generator_{datetime.now().strftime('%Y-%m-%d')}.log"

    body = f"""ANC Sheet generation failed for {date.strftime('%A, %B %d, %Y')}.

Error:
{error}

Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

Log file: {log_file}

Please check the logs and retry manually if needed:
  python3 anc_generator.py {date.strftime('%Y-%m-%d')}

This is an automated notification from the ANC Generator system.
"""

    try:
        # Import email sender from email_utils
        import sys
        email_utils_path = Path(__file__).parent / "email_utils"
        sys.path.insert(0, str(email_utils_path))
        from email_sender import send_email

        success = send_email(
            to=recipients,
            subject=subject,
            body=body,
            method='outlook'
        )

        if success:
            logger.info(f"Failure notification sent to {', '.join(recipients)}")
        else:
            logger.warning("Failed to send failure notification email")

        return success

    except ImportError:
        logger.warning("email_sender module not available for failure notification")
        return False
    except Exception as e:
        logger.error(f"Error sending failure notification: {e}")
        return False

from docx import Document
from docx.shared import Inches, Pt, Twips, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Default font for entire document
DEFAULT_FONT = 'Calibri'

# =============================================================================
# CONFIGURATION LOADING
# =============================================================================

def load_anc_config() -> dict:
    """Load configuration from anc_config.yaml file."""
    config_path = Path(__file__).parent / "anc_config.yaml"
    if config_path.exists():
        logger.debug(f"Loading config from {config_path}")
        with open(config_path) as f:
            config = yaml.safe_load(f)
        logger.info("Configuration loaded successfully")
        return config
    else:
        logger.warning(f"Config file not found at {config_path}, using defaults")
        return {}

# Load config at module level
ANC_CONFIG = load_anc_config()

# =============================================================================
# CONFIG VALIDATION
# =============================================================================

class ConfigValidationError(Exception):
    """Raised when config validation fails."""
    pass

def validate_config(config: dict) -> tuple[bool, list[str]]:
    """
    Validate the ANC configuration file.

    Returns:
        Tuple of (is_valid, list of error/warning messages)
    """
    errors = []
    warnings = []

    # Required sections
    required_sections = ['teams', 'teaching_team_floors']
    for section in required_sections:
        if section not in config:
            errors.append(f"Missing required section: {section}")

    # Validate teams
    teams = config.get('teams', {})
    valid_teams = set(teams.keys()) if teams else set('ABCDEFGHIJ')

    # Validate day configurations
    day_names = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    for day in day_names:
        day_config = config.get(day, {})

        if not day_config:
            warnings.append(f"No configuration for {day}, will use defaults")
            continue

        # Check day_order has valid teams
        day_order = day_config.get('day_order', [])
        for team in day_order:
            if team != 'T' and team not in valid_teams:
                errors.append(f"{day}: Unknown team '{team}' in day_order (valid teams: {valid_teams})")

        # Check evening_order has valid teams
        evening_order = day_config.get('evening_order', [])
        for team in evening_order:
            if team != 'T' and team not in valid_teams:
                errors.append(f"{day}: Unknown team '{team}' in evening_order")

        # Check for duplicate teams in same position (warning only)
        if len(day_order) != len(set(day_order)):
            # This is actually OK - teams can appear multiple times
            pass

    # Validate teaching_team_floors matches teams
    teaching_floors = config.get('teaching_team_floors', {})
    for team in valid_teams:
        if team not in teaching_floors:
            warnings.append(f"No floor assignment for team {team}")

    # Validate holidays format if present
    holidays = config.get('holidays') or {}
    for date_str, holiday_config in holidays.items():
        try:
            datetime.strptime(date_str, '%Y-%m-%d')
        except ValueError:
            errors.append(f"Invalid holiday date format: {date_str} (use YYYY-MM-DD)")

    # Validate overrides format if present
    overrides = config.get('overrides') or {}
    for date_str, override_config in overrides.items():
        try:
            datetime.strptime(date_str, '%Y-%m-%d')
        except ValueError:
            errors.append(f"Invalid override date format: {date_str} (use YYYY-MM-DD)")

    is_valid = len(errors) == 0

    # Log results
    if errors:
        for error in errors:
            logger.error(f"Config validation error: {error}")
    if warnings:
        for warning in warnings:
            logger.warning(f"Config validation warning: {warning}")

    if is_valid:
        logger.info("Config validation passed")
    else:
        logger.error(f"Config validation failed with {len(errors)} error(s)")

    return is_valid, errors + warnings

# =============================================================================
# HOLIDAY AND EXCEPTION HANDLING
# =============================================================================

def is_holiday(date: datetime, config: dict = None) -> tuple[bool, str]:
    """
    Check if a date is a holiday.

    Returns:
        Tuple of (is_holiday, holiday_name or '')
    """
    config = config or ANC_CONFIG
    holidays = config.get('holidays') or {}

    date_str = date.strftime('%Y-%m-%d')
    if date_str in holidays:
        holiday_info = holidays[date_str]
        if isinstance(holiday_info, dict):
            return True, holiday_info.get('name', 'Holiday')
        elif isinstance(holiday_info, str):
            return True, holiday_info
        else:
            return True, 'Holiday'

    return False, ''

def get_date_override(date: datetime, config: dict = None) -> dict:
    """
    Get any pattern override for a specific date.

    Returns:
        Override config dict or empty dict if no override
    """
    config = config or ANC_CONFIG
    overrides = config.get('overrides') or {}

    date_str = date.strftime('%Y-%m-%d')
    if date_str in overrides:
        logger.info(f"Using override configuration for {date_str}")
        return overrides[date_str]

    return {}

def get_effective_day_config(date: datetime, config: dict = None) -> dict:
    """
    Get the effective configuration for a date, considering overrides.

    Priority:
    1. Date-specific override (overrides section)
    2. Holiday configuration (holidays section)
    3. Day-of-week configuration (Monday, Tuesday, etc.)

    Returns:
        Merged configuration dict for the date
    """
    config = config or ANC_CONFIG

    # Get base day config
    day_names = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    day_name = day_names[date.weekday()]
    base_config = config.get(day_name, {}).copy()

    # Check for holiday
    holiday, holiday_name = is_holiday(date, config)
    if holiday:
        holidays = config.get('holidays') or {}
        holiday_config = holidays.get(date.strftime('%Y-%m-%d'), {})
        if isinstance(holiday_config, dict) and holiday_config.get('skip', False):
            logger.info(f"Skipping generation for holiday: {holiday_name}")
            return {'skip': True, 'reason': holiday_name}
        elif isinstance(holiday_config, dict):
            # Merge holiday config with base
            base_config.update(holiday_config)
            base_config['is_holiday'] = True
            base_config['holiday_name'] = holiday_name

    # Check for date-specific override (highest priority)
    override = get_date_override(date, config)
    if override:
        base_config.update(override)
        base_config['has_override'] = True

    return base_config

# =============================================================================
# CONFIGURATION (from YAML or defaults)
# =============================================================================

AMION_ATTENDING_PASSWORD = "tuhs test"
AMION_RESIDENT_PASSWORD = "squirrel"
AMION_BASE_URL = "https://www.amion.com/cgi-bin/ocs"

# Team definitions (from config or defaults)
TEACHING_TEAMS = list(ANC_CONFIG.get('teams', {}).keys()) or ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J']
TEACHING_TEAM_NAMES = ANC_CONFIG.get('teams', {}) or {
    'A': 'Alpha', 'B': 'Bravo', 'C': 'Charlie', 'D': 'Delta', 'E': 'Easy',
    'F': 'Fox', 'G': 'Golf', 'H': 'Hotel', 'I': 'India', 'J': 'Juliet'
}

# Geographic assignments - built from config
def build_team_geography() -> dict:
    """Build team geography mapping from config."""
    geo = {}

    # Teaching team floors from config
    teaching_floors = ANC_CONFIG.get('teaching_team_floors', {})
    if teaching_floors:
        for team, floor in teaching_floors.items():
            geo[team] = floor.split(',')
    else:
        # Default teaching team geography if not in config
        geo.update({
            'A': ['5W', '5E'], 'B': ['5W', '5E'], 'C': ['5W', '5E'], 'D': ['5W', '5E'],
            'E': ['7E', '7W', '8E'], 'F': ['7E', '7W', '8E'],
            'G': ['7E', '7W', '8E'], 'H': ['7E', '7W', '8E'],
            'I': ['3E', '4E'], 'J': ['3E', '4E'],
        })

    # Direct care team floors from config
    direct_care_floors = ANC_CONFIG.get('direct_care_floors', {})
    if direct_care_floors:
        for team, floor in direct_care_floors.items():
            geo[team] = floor.split(',')

    # Also add Med T teams from the list format (backwards compatibility)
    for team in ANC_CONFIG.get('med_t_teams', []):
        name = team.get('name', '')
        floor = team.get('floor', '')
        if name and floor:
            geo[name] = floor.split(',')

    return geo

TEAM_GEOGRAPHY = build_team_geography()

# BAT Phone number (from config or default)
BAT_PHONE = ANC_CONFIG.get('contacts', {}).get('bat_phone', "267-822-6638")

# Manual phone number overrides (from config or defaults)
PHONE_OVERRIDES = ANC_CONFIG.get('phone_overrides', {}) or {
    'Sokach, Carly': '215-498-5369',
    'Tragesser, Lauren': '215-730-8146',
    'Sims, Jason': '215-510-5810',
    'Chau, Alex': '267-908-2851',
}

# =============================================================================
# DATA STRUCTURES
# =============================================================================

@dataclass
class StaffMember:
    name: str
    phone: str = ""
    staff_type: str = ""
    unique_id: str = ""

@dataclass
class TeamAssignment:
    team_letter: str
    team_name: str
    resident: Optional[StaffMember] = None
    attending: Optional[StaffMember] = None
    census: int = 0
    geography: list = field(default_factory=list)

@dataclass
class DirectCareTeam:
    team_number: int
    hospitalist: Optional[StaffMember] = None
    geography: list = field(default_factory=list)

@dataclass
class NightCoverage:
    role: str
    staff: Optional[StaffMember] = None

@dataclass
class AdmissionSlot:
    number: int
    team: str
    geography_hint: str = ""
    time_note: str = ""

# =============================================================================
# AMION API FUNCTIONS
# =============================================================================

def extract_pager_number(pager_str: str) -> str:
    """Extract phone number from pager email format like '2678443524*@paging.templehealth.org'."""
    if not pager_str or '@' not in pager_str:
        return ''
    # Extract digits before the @
    num_part = pager_str.split('@')[0].replace('*', '')
    # Format as phone number if it's 10 digits
    if len(num_part) == 10 and num_part.isdigit():
        return f"{num_part[:3]}-{num_part[3:6]}-{num_part[6:]}"
    return ''

@retry_on_failure(max_retries=3, delay=2.0, backoff=2.0)
def fetch_contact_info(password: str, date: datetime) -> dict[str, str]:
    """Fetch contact info from Amion report 705 and return name->phone mapping."""

    params = {
        'Lo': password,
        'Rpt': '705',
        'Month': date.month,
        'Day': date.day,
        'Year': date.year
    }

    url = f"{AMION_BASE_URL}?{urllib.parse.urlencode(params)}"
    logger.debug(f"Fetching contact info from Amion (Report 705) for {date.strftime('%Y-%m-%d')}")

    with urllib.request.urlopen(url, timeout=30) as response:
        content = response.read().decode('utf-8')

    contacts = {}
    lines = content.strip().split('\n')

    # Skip header lines, parse tab-separated data
    for line in lines[2:]:  # Skip first 2 header lines
        if not line.strip():
            continue

        fields = line.split('\t')
        if len(fields) >= 3:
            # Format: Staff type, Name, Pager, Tel, Email
            name = fields[1].strip()
            pager = fields[2].strip() if len(fields) > 2 else ''
            phone = fields[3].strip() if len(fields) > 3 else ''

            # Use phone if available, otherwise try to extract from pager
            contact = phone if phone else extract_pager_number(pager)

            if name and contact:
                contacts[name] = contact

    logger.debug(f"Retrieved {len(contacts)} contacts from Amion")
    return contacts

@retry_on_failure(max_retries=3, delay=2.0, backoff=2.0)
def fetch_amion_data(password: str, date: datetime, use_academic_year: bool = False) -> list[dict]:
    """Fetch schedule data from Amion API and parse into records."""

    # Amion can use academic years (July-June) - resident schedules typically do
    if use_academic_year:
        # Academic year: July 2025 - June 2026 is "2025"
        year = date.year - 1 if date.month < 7 else date.year
    else:
        year = date.year

    params = {
        'Lo': password,
        'Rpt': '625tabs',
        'Month': date.month,
        'Day': date.day,
        'Year': year
    }

    url = f"{AMION_BASE_URL}?{urllib.parse.urlencode(params)}"
    schedule_type = "resident" if use_academic_year else "attending"
    logger.debug(f"Fetching {schedule_type} schedule from Amion (Report 625tabs) for {date.strftime('%Y-%m-%d')}")

    with urllib.request.urlopen(url, timeout=30) as response:
        content = response.read().decode('utf-8')

    # Check for error responses
    if "NOFI=No file" in content or "no 20" in content.lower():
        logger.warning(f"No schedule file found for {date.strftime('%Y-%m-%d')} ({schedule_type})")
        return []

    records = []
    lines = content.strip().split('\n')

    # Skip header lines (first 6 lines are metadata)
    for line in lines[6:]:
        if not line.strip():
            continue

        fields = line.split('\t')
        if len(fields) >= 10:
            record = {
                'name': fields[0].strip(),
                'unique_id': fields[1].strip() if len(fields) > 1 else '',
                'backup_id': fields[2].strip() if len(fields) > 2 else '',
                'assignment': fields[3].strip() if len(fields) > 3 else '',
                'assignment_id': fields[4].strip() if len(fields) > 4 else '',
                'assignment_backup_id': fields[5].strip() if len(fields) > 5 else '',
                'date': fields[6].strip() if len(fields) > 6 else '',
                'start_time': fields[7].strip() if len(fields) > 7 else '',
                'end_time': fields[8].strip() if len(fields) > 8 else '',
                'staff_type': fields[9].strip() if len(fields) > 9 else '',
                'phone': fields[10].strip() if len(fields) > 10 else '',
            }
            records.append(record)

    return records

def merge_contact_info(records: list[dict], contacts: dict[str, str]) -> list[dict]:
    """Merge contact info into records where phone is missing."""
    for record in records:
        if not record['phone']:
            # Check manual overrides first, then API contacts
            if record['name'] in PHONE_OVERRIDES:
                record['phone'] = PHONE_OVERRIDES[record['name']]
            elif record['name'] in contacts:
                record['phone'] = contacts[record['name']]
    return records

def parse_teaching_teams(attending_data: list[dict], resident_data: list[dict]) -> dict[str, TeamAssignment]:
    """Parse Amion data into teaching team assignments."""

    teams = {}

    for letter in TEACHING_TEAMS:
        team_name = TEACHING_TEAM_NAMES[letter]
        teams[letter] = TeamAssignment(
            team_letter=letter,
            team_name=team_name,
            geography=TEAM_GEOGRAPHY.get(letter, [])
        )

    # Map attending assignments
    attending_patterns = {
        'A': 'Med Alpha', 'B': 'Med Bravo', 'C': 'Med Charlie', 'D': 'Med Delta',
        'E': 'Med Easy', 'F': 'Med Foxtrot', 'G': 'Med Golf', 'H': 'Med Hotel',
        'I': 'Med India', 'J': 'Med Juliet'
    }

    for record in attending_data:
        assignment = record['assignment']
        for letter, pattern in attending_patterns.items():
            if assignment == pattern:
                teams[letter].attending = StaffMember(
                    name=record['name'],
                    phone=record['phone'],
                    staff_type=record['staff_type'],
                    unique_id=record['unique_id']
                )
                break

    # Map resident assignments - look for "Med X Resident" patterns
    for record in resident_data:
        assignment = record['assignment']
        for letter in TEACHING_TEAMS:
            # Match patterns like "Med A Resident - Post", "Med B Resident - Day Admit"
            if f"Med {letter} Resident" in assignment:
                teams[letter].resident = StaffMember(
                    name=record['name'],
                    phone=record['phone'],
                    staff_type=record['staff_type'],
                    unique_id=record['unique_id']
                )
                break

    return teams

def parse_direct_care_teams(attending_data: list[dict]) -> dict[int, DirectCareTeam]:
    """Parse Amion data into direct care team assignments."""

    teams = {}

    for i in range(1, 14):  # Med 1 through Med 13
        teams[i] = DirectCareTeam(
            team_number=i,
            geography=TEAM_GEOGRAPHY.get(f'Med {i}', [])
        )

    for record in attending_data:
        assignment = record['assignment']
        # Match "Med 1", "Med 2", etc. but not "Med Alpha" etc.
        match = re.match(r'^Med (\d+)(?:\s|$)', assignment)
        if match:
            team_num = int(match.group(1))
            if team_num in teams:
                teams[team_num].hospitalist = StaffMember(
                    name=record['name'],
                    phone=record['phone'],
                    staff_type=record['staff_type'],
                    unique_id=record['unique_id']
                )

    return teams

def parse_night_coverage(attending_data: list[dict], resident_data: list[dict]) -> list[NightCoverage]:
    """Parse night coverage assignments."""

    coverage = []

    # Night attending patterns to look for
    night_patterns = [
        ('Temple Nights (Med Q, 1-3)', 'Med 1-3 / IMCU'),
        ('Temple Nights (Med S, 4-6)', 'Med 4-6'),
        ('Temple Nights 3 (Med Y, 7-9)', 'Med 7-9'),
        ('Temple Nights 4 (Med Z, 10-13)', 'Med 10-13'),
    ]

    for record in attending_data:
        assignment = record['assignment']
        for pattern, role in night_patterns:
            if pattern in assignment:
                coverage.append(NightCoverage(
                    role=role,
                    staff=StaffMember(
                        name=record['name'],
                        phone=record['phone'],
                        staff_type=record['staff_type']
                    )
                ))
                break

    # Teaching team night attendings
    teaching_night_patterns = [
        ('Overnight Call A - D, I & J', 'Med A-D, I & J'),
        ('Overnight Call E - H', 'Med E-H'),
    ]

    for record in attending_data:
        assignment = record['assignment']
        for pattern, role in teaching_night_patterns:
            if pattern in assignment:
                coverage.append(NightCoverage(
                    role=role,
                    staff=StaffMember(
                        name=record['name'],
                        phone=record['phone'],
                        staff_type=record['staff_type']
                    )
                ))
                break

    # Resident night assignments - using actual Amion assignment names
    resident_night_patterns = [
        ('Heart Failure & Yellow Nights', 'Heart Failure & Yellow Night Owl'),
        ('AD Night Admitter', 'A-D Resident Admitter'),
        ('EH Night Admitter', 'E-H Resident Admitter'),
        ('ABCD Res Night Float', 'A-D Resident Float'),
        ('EFGH Res Night Float', 'E-H Resident Float'),
        ('N Blue IJ', 'I & J / Med Blue Resident'),
    ]

    for record in resident_data:
        assignment = record['assignment']
        for pattern, role in resident_night_patterns:
            if pattern.lower() in assignment.lower():
                coverage.append(NightCoverage(
                    role=role,
                    staff=StaffMember(
                        name=record['name'],
                        phone=record['phone'],
                        staff_type=record['staff_type']
                    )
                ))
                break

    return coverage

def parse_med_yellow(resident_data: list[dict]) -> list[StaffMember]:
    """Parse Med Yellow team assignments."""

    yellow_staff = []

    for record in resident_data:
        assignment = record['assignment']
        if 'yellow' in assignment.lower() or 'heart failure' in assignment.lower():
            yellow_staff.append(StaffMember(
                name=record['name'],
                phone=record['phone'],
                staff_type=record['staff_type']
            ))

    return yellow_staff

# =============================================================================
# ADMISSION ORDER LOGIC
# =============================================================================

def get_weekday_day_order(date: datetime) -> list[AdmissionSlot]:
    """Generate admission order for weekday day coverage based on day of week."""

    slots = []
    slot_num = 1

    # Get day name for config lookup
    day_names = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    weekday = date.weekday()
    day_name = day_names[weekday]

    # Get day config
    day_config = ANC_CONFIG.get(day_name, {})
    default_floors = ANC_CONFIG.get('default_floors', {})
    day_floors = day_config.get('day_floors', {})

    # Get explicit day_order from config
    day_order = day_config.get('day_order', [])

    if day_order:
        # Use explicit order from config
        for team in day_order:
            if team == 'T':
                slots.append(AdmissionSlot(slot_num, 'T [BAT]', ''))
            else:
                # Get floor from day-specific overrides, then default_floors, then TEAM_GEOGRAPHY
                if team in day_floors:
                    geo = day_floors[team]
                elif team in default_floors:
                    geo = default_floors[team]
                else:
                    geo = ','.join(TEAM_GEOGRAPHY.get(team, []))
                slots.append(AdmissionSlot(slot_num, team, geo))
            slot_num += 1
    else:
        # Fallback to old pattern logic
        default_patterns = {
            0: ['D', 'H', 'A', 'E', 'C', 'G', 'B', 'F'],
            1: ['C', 'G', 'A', 'E', 'I'],
            2: ['B', 'F', 'D', 'H', 'J'],
            3: ['A', 'C', 'E', 'G', 'I'],
            4: ['D', 'H', 'J', 'B', 'F'],
        }
        teaching_rotation = default_patterns.get(weekday, ['A', 'C', 'E', 'G', 'I'])
        for _ in range(2):
            for team in teaching_rotation:
                geo = ','.join(TEAM_GEOGRAPHY.get(team, []))
                slots.append(AdmissionSlot(slot_num, team, geo))
                slot_num += 1

    # Fill remaining with T [BAT] (up to 24 total rows)
    while len(slots) < 24:
        slots.append(AdmissionSlot(slot_num, 'T [BAT]', ''))
        slot_num += 1

    return slots

def get_weekday_evening_order(date: datetime, start_num: int = 23) -> list[AdmissionSlot]:
    """Generate admission order for weekday evening coverage based on day of week."""

    slots = []
    slot_num = start_num

    # Get day name for config lookup
    day_names = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    weekday = date.weekday()
    day_name = day_names[weekday]

    # Get day config
    day_config = ANC_CONFIG.get(day_name, {})
    default_floors = ANC_CONFIG.get('default_floors', {})
    evening_floors = day_config.get('evening_floors', {})
    admission_settings = ANC_CONFIG.get('admission_settings', {})

    # Get time notes from config
    cutoff_note = admission_settings.get('evening_cutoff_note', 'Until 5:30')
    late_time = admission_settings.get('weekday_late_time', '6:30 – 7:00')

    # Get explicit evening_order from config
    evening_order = day_config.get('evening_order', [])
    first_team_rounds = day_config.get('evening_first_team_rounds', 3)

    if evening_order:
        # Filter out 'T' entries to get teaching teams
        teaching_teams = [t for t in evening_order if t != 'T']

        if teaching_teams:
            first_team = teaching_teams[0]
            other_teams = teaching_teams[1:]

            # Build the pattern based on config
            first_team_count = 0

            # Determine how many rounds to do
            for round_num in range(first_team_rounds):
                # First team
                if first_team in evening_floors:
                    geo = evening_floors[first_team]
                elif first_team in default_floors:
                    geo = default_floors[first_team]
                else:
                    geo = ','.join(TEAM_GEOGRAPHY.get(first_team, []))

                first_team_count += 1
                time_note = cutoff_note if first_team_count == first_team_rounds and first_team_rounds > 1 else ""
                slots.append(AdmissionSlot(slot_num, first_team, geo, time_note))
                slot_num += 1

                # Other teams
                for team in other_teams:
                    if team in evening_floors:
                        geo = evening_floors[team]
                    elif team in default_floors:
                        geo = default_floors[team]
                    else:
                        geo = ','.join(TEAM_GEOGRAPHY.get(team, []))
                    slots.append(AdmissionSlot(slot_num, team, geo))
                    slot_num += 1

            # After first team stops, continue with just other teams (2 more)
            if first_team_rounds > 1:
                for team in other_teams:
                    if team in evening_floors:
                        geo = evening_floors[team]
                    elif team in default_floors:
                        geo = default_floors[team]
                    else:
                        geo = ','.join(TEAM_GEOGRAPHY.get(team, []))
                    slots.append(AdmissionSlot(slot_num, team, geo))
                    slot_num += 1
        else:
            # No teaching teams, just use the order as-is (for weekends with T's)
            for team in evening_order:
                if team == 'T':
                    slots.append(AdmissionSlot(slot_num, 'T [BAT]', ''))
                else:
                    if team in evening_floors:
                        geo = evening_floors[team]
                    elif team in default_floors:
                        geo = default_floors[team]
                    else:
                        geo = ','.join(TEAM_GEOGRAPHY.get(team, []))
                    slots.append(AdmissionSlot(slot_num, team, geo))
                slot_num += 1
    else:
        # Fallback to old pattern logic
        default_patterns = {
            0: ['G', 'B', 'F'],
            1: ['J', 'D', 'H'],
            2: ['I', 'C', 'G'],
            3: ['J', 'B', 'F'],
            4: ['I', 'A', 'E'],
        }
        rotation = default_patterns.get(weekday, ['J', 'B', 'F'])
        first_team = rotation[0]
        other_teams = rotation[1:]
        first_team_count = 0

        for round_num in range(3):
            geo = ','.join(TEAM_GEOGRAPHY.get(first_team, []))
            first_team_count += 1
            time_note = cutoff_note if first_team_count == 3 else ""
            slots.append(AdmissionSlot(slot_num, first_team, geo, time_note))
            slot_num += 1

            for team in other_teams:
                geo = ','.join(TEAM_GEOGRAPHY.get(team, []))
                slots.append(AdmissionSlot(slot_num, team, geo))
                slot_num += 1

        for team in other_teams:
            geo = ','.join(TEAM_GEOGRAPHY.get(team, []))
            slots.append(AdmissionSlot(slot_num, team, geo))
            slot_num += 1

    # T [BAT] rows to fill remaining (total should be ~24)
    while len(slots) < 24:
        idx = len(slots) - 11
        time_note = late_time if 4 <= idx <= 9 else ""
        slots.append(AdmissionSlot(slot_num, 'T [BAT]', '', time_note))
        slot_num += 1

    return slots

def get_weekday_night_order(start_num: int = 1) -> list[AdmissionSlot]:
    """Generate admission order for weekday night coverage (7 PM - 7 AM)."""

    slots = []
    slot_num = start_num

    # IJ SNO first (2 slots)
    for _ in range(2):
        slots.append(AdmissionSlot(slot_num, 'IJ SNO', '3E,4E'))
        slot_num += 1

    # Pattern matching original: T [BAT] x4, then EH/AD alternating with T [BAT]
    pattern = [
        'T [BAT]', 'T [BAT]', 'T [BAT]', 'T [BAT]',
        'EH Admitter', 'AD Admitter', 'EH Admitter', 'AD Admitter',
        'T [BAT]', 'T [BAT]', 'T [BAT]',
        'EH Admitter', 'AD Admitter', 'EH Admitter', 'AD Admitter',
        'T [BAT]', 'T [BAT]', 'T [BAT]',
        'EH Admitter', 'AD Admitter', 'EH Float', 'AD Float'
    ]

    for team in pattern:
        geo = '7WE,8E' if 'EH' in team else '5W,5E' if 'AD' in team else ''
        slots.append(AdmissionSlot(slot_num, team, geo))
        slot_num += 1

    return slots

def get_weekday_night_order_continuation(start_num: int = 25) -> list[AdmissionSlot]:
    """Generate continuation of night coverage for page 5."""

    slots = []
    slot_num = start_num

    # Continue the pattern - more T [BAT] and admitters/floats
    pattern = [
        'EH Admitter', 'AD Admitter',
        'T [BAT]', 'T [BAT]', 'T [BAT]', 'T [BAT]',
        'EH Float', 'AD Float', 'EH Admitter', 'AD Admitter',
        'T [BAT]', 'T [BAT]', 'T [BAT]', 'T [BAT]',
        'T [BAT]', 'T [BAT]', 'T [BAT]', 'T [BAT]',
        'T [BAT]', 'T [BAT]', 'T [BAT]', 'T [BAT]'
    ]

    for team in pattern:
        geo = '7WE,8E' if 'EH' in team else '5W,5E' if 'AD' in team else ''
        slots.append(AdmissionSlot(slot_num, team, geo))
        slot_num += 1

    return slots

def get_weekend_day_order(date: datetime) -> list[AdmissionSlot]:
    """Generate admission order for weekend day coverage (7 AM - 5:30 PM)."""

    slots = []
    slot_num = 1

    # Get day name for config lookup
    day_names = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    weekday = date.weekday()
    day_name = day_names[weekday]

    # Get day config
    day_config = ANC_CONFIG.get(day_name, {})
    default_floors = ANC_CONFIG.get('default_floors', {})
    day_floors = day_config.get('day_floors', {})

    # Get explicit day_order from config
    day_order = day_config.get('day_order', [])

    if day_order:
        # Use explicit order from config
        for team in day_order:
            if team == 'T':
                slots.append(AdmissionSlot(slot_num, 'T [BAT]', ''))
            else:
                if team in day_floors:
                    geo = day_floors[team]
                elif team in default_floors:
                    geo = default_floors[team]
                else:
                    geo = ','.join(TEAM_GEOGRAPHY.get(team, []))
                slots.append(AdmissionSlot(slot_num, team, geo))
            slot_num += 1
    else:
        # Fallback to old pattern logic
        default_teams = {
            5: ['B', 'D', 'F', 'H', 'J'],
            6: ['A', 'C', 'E', 'G', 'I'],
        }
        teaching_teams = default_teams.get(weekday, ['A', 'C', 'E', 'G', 'I'])

        for _ in range(4):
            for team in teaching_teams:
                geo = ','.join(TEAM_GEOGRAPHY.get(team, []))
                slots.append(AdmissionSlot(slot_num, team, geo))
                slot_num += 1
                slots.append(AdmissionSlot(slot_num, 'T [BAT]', ''))
                slot_num += 1

    return slots


def get_weekend_evening_order(date: datetime, start_num: int = 23) -> list[AdmissionSlot]:
    """Generate admission order for weekend evening coverage."""

    slots = []
    slot_num = start_num

    # Get day name for config lookup
    day_names = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    weekday = date.weekday()
    day_name = day_names[weekday]

    # Get day config
    day_config = ANC_CONFIG.get(day_name, {})
    default_floors = ANC_CONFIG.get('default_floors', {})
    evening_floors = day_config.get('evening_floors', {})
    admission_settings = ANC_CONFIG.get('admission_settings', {})

    # Get time note for weekends
    late_time = admission_settings.get('weekend_late_time', '5:30-6:30')

    # Get explicit evening_order from config
    evening_order = day_config.get('evening_order', [])

    if evening_order:
        # Use explicit order from config
        for team in evening_order:
            if team == 'T':
                slots.append(AdmissionSlot(slot_num, 'T [BAT]', ''))
            else:
                if team in evening_floors:
                    geo = evening_floors[team]
                elif team in default_floors:
                    geo = default_floors[team]
                else:
                    geo = ','.join(TEAM_GEOGRAPHY.get(team, []))
                slots.append(AdmissionSlot(slot_num, team, geo))
            slot_num += 1
    else:
        # Fallback to old pattern logic
        default_teams = {
            5: ['D', 'F', 'H', 'J'],
            6: ['C', 'E', 'G', 'I'],
        }
        evening_teams = default_teams.get(weekday, ['C', 'E', 'G', 'I'])

        for _ in range(3):
            for team in evening_teams:
                geo = ','.join(TEAM_GEOGRAPHY.get(team, []))
                slots.append(AdmissionSlot(slot_num, team, geo))
                slot_num += 1
                slots.append(AdmissionSlot(slot_num, 'T [BAT]', ''))
                slot_num += 1

    # Fill remaining with T [BAT] with late time annotations
    while len(slots) < 24:
        idx = len(slots) - 7
        time_note = late_time if 0 <= idx <= 8 else ""
        slots.append(AdmissionSlot(slot_num, 'T [BAT]', '', time_note))
        slot_num += 1

    return slots

def get_admission_order(date: datetime) -> dict:
    """Get the full admission order for a given date based on day of week."""

    weekday = date.weekday()  # 0=Monday, 6=Sunday

    if weekday < 5:  # Monday-Friday (weekday)
        return {
            'day': get_weekday_day_order(date),
            'evening': get_weekday_evening_order(date),
            'night': get_weekday_night_order(),
            'night_continuation': get_weekday_night_order_continuation(),
            'is_weekend': False,
            'day_coverage_times': {
                'start': '7:30 AM',
                'end': '2:30 PM',
                'protected_time': '11:30 AM - 1 PM'
            },
            'evening_coverage_times': {
                'start': '2:30 PM',
                'end': '6:30 PM'
            }
        }
    else:  # Saturday-Sunday (weekend)
        return {
            'day': get_weekend_day_order(date),
            'evening': get_weekend_evening_order(date),
            'night': get_weekday_night_order(),
            'night_continuation': get_weekday_night_order_continuation(),
            'is_weekend': True,
            'day_coverage_times': {
                'start': '7 AM',
                'end': '5:30 PM',
                'protected_time': None
            }
        }

# =============================================================================
# DOCUMENT GENERATION
# =============================================================================

def generate_anc_sheet(date: datetime) -> dict:
    """Generate all data needed for an ANC sheet."""

    logger.info(f"Generating ANC sheet for {date.strftime('%A, %B %d, %Y')}")

    # Check for holidays or date-specific overrides
    day_config = get_effective_day_config(date)
    if day_config.get('skip'):
        reason = day_config.get('reason', 'Holiday')
        logger.info(f"Skipping generation: {reason}")
        raise ValueError(f"Generation skipped: {reason}")

    if day_config.get('is_holiday'):
        logger.info(f"Holiday: {day_config.get('holiday_name', 'Unknown')}")
    if day_config.get('has_override'):
        logger.info("Using date-specific override configuration")

    # Fetch data from both Amion schedules with retry logic
    # Attending schedule uses calendar year, resident schedule uses academic year
    logger.info("Fetching Amion data...")

    try:
        attending_data = fetch_amion_data(AMION_ATTENDING_PASSWORD, date, use_academic_year=False)
    except Exception as e:
        logger.error(f"Failed to fetch attending data: {e}")
        attending_data = []

    try:
        resident_data = fetch_amion_data(AMION_RESIDENT_PASSWORD, date, use_academic_year=True)
    except Exception as e:
        logger.error(f"Failed to fetch resident data: {e}")
        resident_data = []

    # Fetch contact info (report 705) to fill in missing phone numbers
    try:
        attending_contacts = fetch_contact_info(AMION_ATTENDING_PASSWORD, date)
        attending_data = merge_contact_info(attending_data, attending_contacts)
    except Exception as e:
        logger.warning(f"Failed to fetch contact info, continuing without: {e}")

    if not attending_data:
        logger.warning("No attending data retrieved")
    if not resident_data:
        logger.warning("No resident data retrieved")

    logger.info(f"Retrieved {len(attending_data)} attending assignments, {len(resident_data)} resident assignments")

    # Parse into structured data
    teaching_teams = parse_teaching_teams(attending_data, resident_data)
    direct_care_teams = parse_direct_care_teams(attending_data)
    night_coverage = parse_night_coverage(attending_data, resident_data)
    med_yellow = parse_med_yellow(resident_data)
    admission_order = get_admission_order(date)

    # Parse Physician Executive
    physician_executive = None
    for record in attending_data:
        if record['assignment'] == 'Physician Executive':
            physician_executive = StaffMember(
                name=record['name'],
                phone=record['phone'],
                staff_type=record['staff_type']
            )
            break

    logger.debug("Data parsing complete")

    return {
        'date': date,
        'teaching_teams': teaching_teams,
        'direct_care_teams': direct_care_teams,
        'night_coverage': night_coverage,
        'med_yellow': med_yellow,
        'admission_order': admission_order,
        'bat_phone': BAT_PHONE,
        'physician_executive': physician_executive,
        'day_config': day_config  # Include for reference
    }

# =============================================================================
# WORD DOCUMENT GENERATION
# =============================================================================

def set_cell_shading(cell, color: str):
    """Set background color for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_vertical_center(cell):
    """Set vertical alignment to center for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)

def format_cell(cell, font_size: int = 9, bold: bool = False, shading: str = None, vertical_center: bool = True):
    """Apply standard formatting to a table cell."""
    if vertical_center:
        set_cell_vertical_center(cell)
    if shading:
        set_cell_shading(cell, shading)
    for para in cell.paragraphs:
        # If paragraph has text but no runs, we need to handle it
        if para.runs:
            for run in para.runs:
                run.font.name = DEFAULT_FONT
                run.font.size = Pt(font_size)
                run.bold = bold
        elif para.text:
            # Clear and re-add with formatting
            text = para.text
            para.clear()
            run = para.add_run(text)
            run.font.name = DEFAULT_FONT
            run.font.size = Pt(font_size)
            run.bold = bold

def set_row_height(row, height_twips: int):
    """Set the height of a table row in twips (1/20 of a point)."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height_twips))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def add_section_heading(doc: Document, text: str, font_size: int = 10) -> None:
    """Add a consistent section heading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = DEFAULT_FONT
    run.bold = True
    run.font.size = Pt(font_size)

def create_roster_table(doc: Document, data: dict) -> None:
    """Create the main roster table with teaching teams and direct care."""

    # Row heights (in twips - 1/20 of a point)
    TEACHING_ROW_HEIGHT = 480  # Tall enough for 2 lines of text
    COMPACT_ROW_HEIGHT = 260   # Compact for single-line entries

    # Softer colors for legibility
    SOFT_RED = 'FFCCCB'      # Soft red/pink
    SOFT_BLUE = 'CCE5FF'     # Soft blue
    SOFT_YELLOW = 'FFFFCC'   # Soft yellow

    # Teaching teams table (6 columns) - no title, self-explanatory
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # Header row - SOFT RED background
    headers = ['Team', '7AM Census', 'Resident (admits)', 'Resident Phone',
               'Attending', 'Attending Phone']
    header_row = table.rows[0]
    set_row_height(header_row, COMPACT_ROW_HEIGHT)
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        format_cell(cell, font_size=8, bold=True, shading=SOFT_RED)

    # Teaching team rows (A-J)
    for letter in TEACHING_TEAMS:
        team = data['teaching_teams'][letter]
        row = table.add_row()
        set_row_height(row, TEACHING_ROW_HEIGHT)

        # Team name
        row.cells[0].text = f"{letter} ~ {team.team_name}"

        # Census (blank for manual entry)
        row.cells[1].text = ""

        # Resident info
        if team.resident:
            row.cells[2].text = team.resident.name
            row.cells[3].text = team.resident.phone

        # Attending info
        if team.attending:
            row.cells[4].text = team.attending.name
            row.cells[5].text = team.attending.phone

        # Apply formatting to all cells
        for cell in row.cells:
            format_cell(cell, font_size=9)

    # Med Yellow row
    row = table.add_row()
    set_row_height(row, TEACHING_ROW_HEIGHT)
    row.cells[0].text = "Med Yellow"
    format_cell(row.cells[0], font_size=9, shading=SOFT_YELLOW)

    # Get Med Yellow staff
    yellow_names = [s.name for s in data['med_yellow'][:2]]
    yellow_phones = [s.phone for s in data['med_yellow'][:2]]
    row.cells[2].text = '\n'.join(yellow_names) if yellow_names else ""
    row.cells[3].text = '\n'.join(yellow_phones) if yellow_phones else ""

    # Apply formatting to all Med Yellow cells
    for cell in row.cells:
        format_cell(cell, font_size=9)

    # Direct Care teams table - no title heading
    dc_table = doc.add_table(rows=1, cols=3)
    dc_table.style = 'Table Grid'

    # Header - SOFT BLUE background
    dc_headers = ['Team', 'Hospitalist', 'Phone']
    dc_header_row = dc_table.rows[0]
    set_row_height(dc_header_row, COMPACT_ROW_HEIGHT)
    for i, header in enumerate(dc_headers):
        dc_header_row.cells[i].text = header
        format_cell(dc_header_row.cells[i], font_size=8, bold=True, shading=SOFT_BLUE)

    # BAT Phone row first
    bat_row = dc_table.add_row()
    set_row_height(bat_row, COMPACT_ROW_HEIGHT)
    bat_row.cells[0].text = "BAT Phone"
    bat_row.cells[1].text = ""
    bat_row.cells[2].text = data['bat_phone']
    for cell in bat_row.cells:
        format_cell(cell, font_size=8, bold=True)

    # Physician Executive row
    if data.get('physician_executive'):
        pe_row = dc_table.add_row()
        set_row_height(pe_row, COMPACT_ROW_HEIGHT)
        pe_row.cells[0].text = "Physician Executive"
        pe_row.cells[1].text = data['physician_executive'].name
        pe_row.cells[2].text = data['physician_executive'].phone
        for cell in pe_row.cells:
            format_cell(cell, font_size=8, bold=True)

    # Direct care team rows - compact height
    for num in range(1, 14):
        team = data['direct_care_teams'][num]
        row = dc_table.add_row()
        set_row_height(row, COMPACT_ROW_HEIGHT)
        # Med 1-3 are IMCU teams
        if num <= 3:
            row.cells[0].text = f"Med {num} / IMCU"
        else:
            row.cells[0].text = f"Med {num}"
        if team.hospitalist:
            row.cells[1].text = team.hospitalist.name
            row.cells[2].text = team.hospitalist.phone
        for cell in row.cells:
            format_cell(cell, font_size=8)

def create_night_coverage_table(doc: Document, data: dict) -> None:
    """Create the night coverage section."""

    # Compact row height (same as direct care)
    COMPACT_ROW_HEIGHT = 260

    # Soft purple for legibility
    SOFT_PURPLE = 'E6D5F2'

    # No title - just the table (self-explanatory with colored header)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Header - SOFT PURPLE background
    headers = ['Night Coverage', 'Resident/Attending', 'Contact #']
    header_row = table.rows[0]
    set_row_height(header_row, COMPACT_ROW_HEIGHT)
    for i, header in enumerate(headers):
        header_row.cells[i].text = header
        format_cell(header_row.cells[i], font_size=8, bold=True, shading=SOFT_PURPLE)

    for coverage in data['night_coverage']:
        row = table.add_row()
        set_row_height(row, COMPACT_ROW_HEIGHT)
        row.cells[0].text = coverage.role
        if coverage.staff:
            row.cells[1].text = coverage.staff.name
            row.cells[2].text = coverage.staff.phone
        for cell in row.cells:
            format_cell(cell, font_size=8)

def set_column_widths(table, widths: list):
    """Set column widths for a table (widths in inches)."""
    for row in table.rows:
        for i, width in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(width)

def create_yellow_hf_table(doc: Document) -> None:
    """Create the Yellow/Heart Failure SNO tracking table."""

    COMPACT_ROW_HEIGHT = 240
    ROW_HEIGHT = 340  # Smaller to fit on page with Holds
    SOFT_YELLOW = 'FFFFCC'

    # Day Coverage section
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'

    # Header spanning all columns
    header_row = table.rows[0]
    set_row_height(header_row, COMPACT_ROW_HEIGHT)
    header_row.cells[0].merge(header_row.cells[6])
    header_row.cells[0].text = "Day Coverage 7:00 AM - 6:00 PM"
    format_cell(header_row.cells[0], font_size=9, bold=True, shading=SOFT_YELLOW)

    # Column headers
    col_headers = ['#', 'Team', 'Time', 'Patient Name', 'MRN', 'Origin', 'S/O contact']
    header_row2 = table.add_row()
    set_row_height(header_row2, COMPACT_ROW_HEIGHT)
    for i, header in enumerate(col_headers):
        header_row2.cells[i].text = header
        format_cell(header_row2.cells[i], font_size=8, bold=True)

    # Yellow rows (5 slots)
    for i in range(1, 6):
        row = table.add_row()
        set_row_height(row, ROW_HEIGHT)
        row.cells[0].text = str(i)
        row.cells[1].text = "Yellow"
        for cell in row.cells:
            format_cell(cell, font_size=8)

    # Night Coverage section header
    night_header = table.add_row()
    set_row_height(night_header, COMPACT_ROW_HEIGHT)
    night_header.cells[0].merge(night_header.cells[6])
    night_header.cells[0].text = "Night Coverage 7:00 PM - 6:00 AM"
    format_cell(night_header.cells[0], font_size=9, bold=True, shading=SOFT_YELLOW)

    # Night column headers
    night_col_row = table.add_row()
    set_row_height(night_col_row, COMPACT_ROW_HEIGHT)
    for i, header in enumerate(col_headers):
        night_col_row.cells[i].text = header
        format_cell(night_col_row.cells[i], font_size=8, bold=True)

    # HF/Yellow SNO rows (3 slots)
    for i in range(1, 4):
        row = table.add_row()
        set_row_height(row, ROW_HEIGHT)
        row.cells[0].text = str(i)
        row.cells[1].text = "HF/Yellow SNO"
        for cell in row.cells:
            format_cell(cell, font_size=8)

def create_holds_table(doc: Document) -> None:
    """Create the Holds tracking table for downgrades/RICU/CICU."""

    COMPACT_ROW_HEIGHT = 240
    ROW_HEIGHT = 340  # Smaller to fit on same page
    SOFT_GRAY = 'E0E0E0'

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # Header spanning all columns
    header_row = table.rows[0]
    set_row_height(header_row, COMPACT_ROW_HEIGHT)
    header_row.cells[0].merge(header_row.cells[5])
    header_row.cells[0].text = "Holds - Boyer downgrades/RICU/CICU"
    format_cell(header_row.cells[0], font_size=9, bold=True, shading=SOFT_GRAY)

    # Column headers
    col_headers = ['#', 'Origin', 'Time', 'Patient Name', 'MRN', 'New Team']
    header_row2 = table.add_row()
    set_row_height(header_row2, COMPACT_ROW_HEIGHT)
    for i, header in enumerate(col_headers):
        header_row2.cells[i].text = header
        format_cell(header_row2.cells[i], font_size=8, bold=True)

    # Empty rows for manual entry (17 slots - reduced to fit on page)
    for i in range(1, 18):
        row = table.add_row()
        set_row_height(row, ROW_HEIGHT)
        row.cells[0].text = str(i)
        for cell in row.cells:
            format_cell(cell, font_size=8)

def create_admission_order_table(doc: Document, data: dict, section: str, title: str, rules: str, add_page_break: bool = False) -> None:
    """Create an admission order table for a coverage period."""

    # Taller row height for writing (480 twips = 24pt, good for pencil entry)
    ROW_HEIGHT = 480

    if add_page_break:
        doc.add_page_break()

    add_section_heading(doc, title, 11)

    # Rules text - formatted for legibility
    p = doc.add_paragraph()
    # Split rules into title and body for better formatting
    rules_lines = rules.split('\n')
    if len(rules_lines) > 1:
        # First line is the coverage period description
        run = p.add_run(rules_lines[0])
        run.font.name = DEFAULT_FONT
        run.font.size = Pt(9)
        run.bold = True
        # Add line break and rest of rules
        p.add_run('\n')
        run2 = p.add_run(rules_lines[1])
        run2.font.name = DEFAULT_FONT
        run2.font.size = Pt(8)
    else:
        run = p.add_run(rules)
        run.font.name = DEFAULT_FONT
        run.font.size = Pt(8)

    slots = data['admission_order'].get(section, [])
    if not slots:
        return

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # Set column widths: # much narrower than team
    # Total page width ~7.5" with 0.5" margins = 6.5" usable
    # #=0.25 (very narrow), Team=1.35, Time=0.45, Patient=2.35, MRN=0.8, Origin=1.3
    set_column_widths(table, [0.25, 1.35, 0.45, 2.35, 0.8, 1.3])

    headers = ['#', 'Team', 'Time', 'Patient Name', 'MRN', 'Origin S/O contact']
    header_row = table.rows[0]
    set_row_height(header_row, ROW_HEIGHT)
    for i, header in enumerate(headers):
        header_row.cells[i].text = header
        format_cell(header_row.cells[i], font_size=8, bold=True, shading='D9D9D9')

    for slot in slots:
        row = table.add_row()
        set_row_height(row, ROW_HEIGHT)
        row.cells[0].text = str(slot.number)

        # Format team with geography hint - use spaces: C (5W, 5E)
        team_text = slot.team
        if slot.geography_hint:
            # Add spaces after commas in geography hint
            geo_formatted = slot.geography_hint.replace(',', ', ')
            team_text += f" ({geo_formatted})"
        row.cells[1].text = team_text

        row.cells[2].text = slot.time_note
        # Leave patient name, MRN, origin blank for manual entry

        for cell in row.cells:
            format_cell(cell, font_size=8)

def generate_word_document(data: dict, output_path: str) -> str:
    """Generate a complete ANC sheet Word document."""

    doc = Document()

    # Set default font for document
    doc.styles['Normal'].font.name = DEFAULT_FONT

    # Set narrow margins and add header with date
    date = data['date']
    header_date_str = date.strftime('%A, %B %d, %Y')

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        # Different first page header (no header on first page)
        section.different_first_page_header_footer = True

        # Add header with date (faint gray text) - appears on pages 2+
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.clear()
        header_run = header_para.add_run(header_date_str)
        header_run.font.name = DEFAULT_FONT
        header_run.font.size = Pt(9)
        header_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # First page header is empty (just leave it blank)

    # Title format: "Wednesday, January 14, 2026"
    title_date_str = date.strftime('%A, %B %d, %Y')

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run(title_date_str)
    title_run.font.name = DEFAULT_FONT
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Roster tables
    create_roster_table(doc, data)

    # Night coverage
    create_night_coverage_table(doc, data)

    # Admission order tables
    admission_order = data['admission_order']

    # Get admission rules from config
    rules_config = ANC_CONFIG.get('admission_rules', {})

    if admission_order['is_weekend']:
        # Weekend format - page 2 starts with day coverage
        day_header = rules_config.get('weekend_day_header', "Day Coverage (7 AM - 5:30 PM)")
        day_rules_text = rules_config.get('weekend_day_rules',
            "Teaching admits 7 AM - 5:30 PM. No protected time. "
            "Esperanza to teaching. Obs (chest pain, asthma, COPD, syncope) to Med T. "
            "Once capped, all admissions to Med T. No admissions to capped teams (16 A-J).")
        rules = f"{day_header}\n{day_rules_text.strip()}"
        create_admission_order_table(doc, data, 'day',
                                     f"Day Coverage ({admission_order['day_coverage_times']['start']} - {admission_order['day_coverage_times']['end']})",
                                     rules, add_page_break=True)
    else:
        # Weekday format - page 2 starts with day coverage
        # Check if Thursday (academic half day)
        is_thursday = date.weekday() == 3

        if is_thursday:
            day_header = rules_config.get('thursday_day_header', "Day Rules - Academic Half Day")
            day_rules_text = rules_config.get('thursday_day_rules',
                "Academic half day with protected time from 11:30 AM - 2:30 PM. "
                "Teaching admits 7 - 11:30 AM only. All admissions to Med T during protected time. "
                "Esperanza to teaching. Obs (chest pain, asthma, COPD, syncope) to Med T. "
                "Once capped, all admissions to Med T. No admissions to capped teams (16 A-J).")
        else:
            day_header = rules_config.get('weekday_day_header', "Day Coverage Teaching admits (7:30 – 2:30PM)")
            day_rules_text = rules_config.get('weekday_day_rules',
                "Teaching admits 7 - 11:30 AM, 1 - 2:30 PM. Protected time 11:30 AM - 1 PM. "
                "Esperanza to teaching. Obs (chest pain, asthma, COPD, syncope) to Med T. "
                "Once capped, all admissions to Med T. No admissions to capped teams (16 A-J). "
                "4west, 6 East/West– can go to any team. Direct Care: 3E/4E preferentially.")
        day_rules = f"{day_header}\n{day_rules_text.strip()}"
        create_admission_order_table(doc, data, 'day', "Day Coverage (7:30 AM - 2:30 PM)", day_rules, add_page_break=True)

        evening_header = rules_config.get('weekday_evening_header', "Evening Coverage (2:30 – 6:30pm)")
        evening_rules_text = rules_config.get('weekday_evening_rules',
            "Teaching admits until 6:30PM (A-H). Med I & J admit until 5:30. "
            "Esperanza to teaching. Obs (chest pain, asthma, COPD, syncope) to Med T. "
            "Once capped, all admissions to Med T. No admissions to capped teams (16 A-J). "
            "4west, 6 East/West– can go to any team. Direct Care: 3E/4E preferentially.")
        evening_rules = f"{evening_header}\n{evening_rules_text.strip()}"
        create_admission_order_table(doc, data, 'evening', "Evening Coverage (2:30 PM - 6:30 PM)", evening_rules, add_page_break=True)

    # Night coverage admission order - page 4
    night_header = rules_config.get('weekday_night_header', "Night Coverage 7:00 PM - 7:00 AM")
    night_rules_text = rules_config.get('weekday_night_rules',
        "AD, EH Admitters cannot admit prior to 8 PM: must 'skip ahead' and give to Med T if needed. "
        "Teaching stops at 6:15 AM. Esperanza to teaching. "
        "Obs (chest pain, asthma, COPD, syncope) to Med T.")
    night_rules = f"{night_header}\n{night_rules_text.strip()}"
    create_admission_order_table(doc, data, 'night', "Night Coverage (7:00 PM - 7:00 AM)", night_rules, add_page_break=True)

    # Night coverage continuation - page 5
    night_cont_rules = ("Night Coverage Continued\n"
                       "AD, EH Admitters cannot admit prior to 8 PM: must 'skip ahead' and give to Med T if needed. "
                       "Teaching stops at 6:15 AM.")
    create_admission_order_table(doc, data, 'night_continuation', "Night Coverage (Continued)", night_cont_rules, add_page_break=True)

    # Yellow/HF SNO table and Holds table - page 6 (both on same page)
    doc.add_page_break()

    # Add Subspecialty Coverage header (left-justified)
    subspecialty_header = rules_config.get('subspecialty_header', "Subspecialty Coverage")
    subspec_para = doc.add_paragraph()
    subspec_run = subspec_para.add_run(subspecialty_header)
    subspec_run.font.name = DEFAULT_FONT
    subspec_run.bold = True
    subspec_run.font.size = Pt(12)
    subspec_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    create_yellow_hf_table(doc)
    create_holds_table(doc)

    # Save document
    doc.save(output_path)
    return output_path

def convert_docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """Convert a Word document to PDF.

    Tries docx2pdf (requires MS Word) or LibreOffice headless mode.
    Returns False if no conversion method available.
    """

    import platform
    system = platform.system()

    # Try docx2pdf first (uses Word on macOS/Windows)
    try:
        from docx2pdf import convert
        print("Converting to PDF via docx2pdf...")
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            return True
    except ImportError:
        pass
    except Exception as e:
        print(f"docx2pdf failed: {e}")

    # Try LibreOffice headless (no GUI)
    libreoffice_paths = [
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',
        '/usr/local/bin/soffice',
        'soffice',
        'libreoffice'
    ]

    for path in libreoffice_paths:
        try:
            result = subprocess.run([path, '--version'], capture_output=True, timeout=5)
            if result.returncode == 0:
                print("Converting to PDF via LibreOffice...")
                output_dir = os.path.dirname(pdf_path)
                result = subprocess.run([
                    path, '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', output_dir,
                    docx_path
                ], capture_output=True, timeout=60)
                if result.returncode == 0:
                    expected_pdf = os.path.splitext(docx_path)[0] + '.pdf'
                    if expected_pdf != pdf_path and os.path.exists(expected_pdf):
                        os.rename(expected_pdf, pdf_path)
                    if os.path.exists(pdf_path):
                        return True
        except (subprocess.TimeoutExpired, FileNotFoundError):
            continue

    # No conversion method worked
    print("\nPDF conversion requires LibreOffice (free) or Microsoft Word.")
    print("Install LibreOffice from: https://www.libreoffice.org/download/")
    return False

def generate_anc_for_date(date: datetime, output_dir: str = None, output_format: str = 'pdf',
                         validate: bool = True, notify_on_failure: bool = True) -> str:
    """Generate an ANC sheet for a specific date and save to file.

    Args:
        date: The date to generate the ANC sheet for
        output_dir: Directory to save the file (defaults to script directory)
        output_format: 'pdf' (default) or 'docx'
        validate: Whether to validate config before generation
        notify_on_failure: Whether to send email notification on failure

    Returns:
        Path to the generated file

    Raises:
        ConfigValidationError: If config validation fails
        Exception: If generation fails
    """
    logger.info(f"Starting ANC generation for {date.strftime('%Y-%m-%d')}")

    # Validate config if requested
    if validate:
        is_valid, messages = validate_config(ANC_CONFIG)
        if not is_valid:
            error_msg = f"Config validation failed: {'; '.join(messages)}"
            logger.error(error_msg)
            if notify_on_failure:
                send_failure_notification(date, error_msg)
            raise ConfigValidationError(error_msg)

    if output_dir is None:
        output_dir = os.path.dirname(os.path.abspath(__file__))

    # Generate filename with AUTO_ prefix to indicate auto-generated
    date_str = date.strftime('%m %d %y %A %B %Y')

    try:
        # Generate data
        data = generate_anc_sheet(date)

        if output_format.lower() == 'pdf':
            # Create Word doc in temp location, convert to PDF
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                temp_docx = tmp.name

            try:
                generate_word_document(data, temp_docx)

                pdf_filename = f"AUTO_{date_str} ANC Sheets.pdf"
                pdf_path = os.path.join(output_dir, pdf_filename)

                if convert_docx_to_pdf(temp_docx, pdf_path):
                    logger.info(f"Generated PDF: {pdf_path}")
                    return pdf_path
                else:
                    # Fallback to docx if PDF conversion fails
                    logger.warning("PDF conversion failed, saving as Word document instead")
                    docx_filename = f"AUTO_{date_str} ANC Sheets.docx"
                    docx_path = os.path.join(output_dir, docx_filename)
                    os.rename(temp_docx, docx_path)
                    logger.info(f"Generated DOCX: {docx_path}")
                    return docx_path
            finally:
                # Clean up temp file if it still exists
                if os.path.exists(temp_docx):
                    os.remove(temp_docx)
        else:
            # Output as Word document
            docx_filename = f"AUTO_{date_str} ANC Sheets.docx"
            docx_path = os.path.join(output_dir, docx_filename)
            generate_word_document(data, docx_path)
            logger.info(f"Generated DOCX: {docx_path}")
            return docx_path

    except Exception as e:
        error_msg = f"Generation failed: {str(e)}"
        logger.error(error_msg, exc_info=True)
        if notify_on_failure:
            send_failure_notification(date, error_msg)
        raise

# =============================================================================
# MAIN / TEST
# =============================================================================

def prompt_with_timeout(prompt: str, timeout: int = 7, default: str = '1') -> str:
    """Prompt user for input with a timeout. Returns default if no input received."""
    import select
    import sys

    print(prompt, end='', flush=True)

    # Use select to wait for input with timeout (Unix/macOS only)
    ready, _, _ = select.select([sys.stdin], [], [], timeout)

    if ready:
        response = sys.stdin.readline().strip()
        return response if response else default
    else:
        print(f"\n(No input received, defaulting to {default})")
        return default

if __name__ == "__main__":
    import sys

    print("=" * 60)
    print("ANC Sheet Generator")
    print("=" * 60)

    # Parse command line arguments for date
    target_date = None
    format_from_arg = None

    for arg in sys.argv[1:]:
        if arg.lower() in ('--docx', '-d'):
            format_from_arg = 'docx'
        elif arg.lower() in ('--pdf', '-p'):
            format_from_arg = 'pdf'
        elif target_date is None:
            # Date provided as argument (format: MM-DD-YYYY or YYYY-MM-DD)
            try:
                if '-' in arg and len(arg.split('-')[0]) == 4:
                    target_date = datetime.strptime(arg, '%Y-%m-%d')
                else:
                    target_date = datetime.strptime(arg, '%m-%d-%Y')
            except ValueError:
                print(f"Invalid date format: {arg}")
                print("Use: MM-DD-YYYY or YYYY-MM-DD")
                print("Options: --pdf, --docx")
                sys.exit(1)

    if target_date is None:
        # Default to today
        target_date = datetime.now()

    print(f"Date: {target_date.strftime('%A, %B %d, %Y')}")
    print()

    # If format not specified via command line, prompt user
    if format_from_arg:
        output_format = format_from_arg
        print(f"Output format: {output_format.upper()}")
    else:
        print("Select output format:")
        print("  1. Word (.docx) - default")
        print("  2. PDF")
        print()
        choice = prompt_with_timeout("Enter choice (1 or 2) [7 sec timeout]: ", timeout=7, default='1')

        if choice == '2':
            output_format = 'pdf'
        else:
            output_format = 'docx'

        print(f"Output format: {output_format.upper()}")

    print()

    # Generate the document
    output_path = generate_anc_for_date(target_date, output_format=output_format)

    print()
    print("=" * 60)
    print("Generation complete!")
    print(f"Output: {output_path}")
    print("=" * 60)
