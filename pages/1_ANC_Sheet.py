#!/usr/bin/env python3
"""
ANC Sheet Viewer - View the current ANC sheet
"""

import streamlit as st
from pathlib import Path
from datetime import datetime, date
import re

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

st.set_page_config(
    page_title="ANC Sheet - Geo Owl",
    page_icon="Gemini_Generated_Image_2hkaog2hkaog2hka.png",
    layout="wide"
)

st.title("ANC Sheet")

# Date picker (default to today)
selected_date = st.date_input("Select date", value=date.today())

# Find ANC sheet for selected date
project_dir = Path(__file__).parent.parent

# Build expected filename pattern: AUTO_MM DD YY
date_str = selected_date.strftime("%m %d %y")
pattern = f"AUTO_{date_str}*.docx"
matching_files = list(project_dir.glob(pattern))

if matching_files:
    anc_file = matching_files[0]

    # Download button
    col1, col2 = st.columns([3, 1])
    with col1:
        st.success(f"Found: **{anc_file.name}**")
    with col2:
        with open(anc_file, "rb") as f:
            st.download_button(
                "Download",
                data=f.read(),
                file_name=anc_file.name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # Preview content
    st.markdown("---")
    st.subheader("Preview")

    if DOCX_AVAILABLE:
        try:
            doc = Document(anc_file)

            # Extract tables (ANC sheets are mostly tables)
            for i, table in enumerate(doc.tables):
                if i >= 1:  # Only show first table (first page)
                    break

                # Convert table to markdown-ish format
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                if table_data:
                    # Display as dataframe for better formatting
                    import pandas as pd
                    df = pd.DataFrame(table_data[1:], columns=table_data[0] if table_data else None)
                    st.dataframe(df, use_container_width=True, hide_index=True)

            # Also show any paragraphs (headers, etc.)
            for para in doc.paragraphs[:10]:  # First 10 paragraphs
                text = para.text.strip()
                if text:
                    st.markdown(text)

        except Exception as e:
            st.error(f"Error reading document: {e}")
            st.info("Download the file to view the full content.")
    else:
        st.warning("Install python-docx to enable preview: `pip install python-docx`")
        st.info("Download the file to view the full content.")

else:
    st.warning(f"No ANC sheet found for {selected_date.strftime('%A, %B %d, %Y')}")
    st.markdown("""
    **To generate an ANC sheet**, run from terminal:
    ```
    python anc_generator.py
    ```
    """)

    # Show available dates
    all_files = list(project_dir.glob("AUTO_*.docx"))
    if all_files:
        st.markdown("---")
        st.subheader("Available dates")
        for f in sorted(all_files, reverse=True)[:5]:
            # Parse date from filename: AUTO_MM DD YY ...
            match = re.search(r'AUTO_(\d{2}) (\d{2}) (\d{2})', f.name)
            if match:
                try:
                    file_date = datetime.strptime(f"{match.group(1)}/{match.group(2)}/{match.group(3)}", "%m/%d/%y")
                    st.caption(f"- {file_date.strftime('%A, %B %d, %Y')}")
                except:
                    st.caption(f"- {f.name}")
